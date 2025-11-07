import pytesseract
import pdfplumber
from PIL import Image
import cv2
import numpy as np
import ollama
import base64
import io
import os
import re
from docx import Document
import tempfile
import pypdfium2 as pdfium

# Ensure Tesseract is discoverable on Windows by setting a sensible default path
try:
    # This will succeed if tesseract is on PATH or already configured
    _ = pytesseract.get_tesseract_version()
except Exception:
    # Try common Windows install paths or env override
    possible_paths = [
        os.environ.get("TESSERACT_PATH"),
        r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe",
        r"C:\\Program Files (x86)\\Tesseract-OCR\\tesseract.exe",
    ]
    for candidate in possible_paths:
        if candidate and os.path.exists(candidate):
            pytesseract.pytesseract.tesseract_cmd = candidate
            break

# Function to extract text from PDFs using pdfplumber
def extract_text_from_pdf(file_path, max_pages_for_vision: int = 5, render_scale: float = 2.0):
    """Extract text from PDF and augment with per-page visual analysis.

    - Uses pdfplumber to extract selectable text.
    - Renders up to `max_pages_for_vision` pages with pypdfium2 and analyzes them with the vision model
      to capture diagrams, tables, and non-selectable (scanned) text.
    """
    text_sections = []

    # 1) Extract selectable text
    try:
        with pdfplumber.open(file_path) as pdf:
            extracted = []
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    extracted.append(page_text)
            if extracted:
                text_sections.append("Selectable text from PDF:\n" + "\n\n".join(extracted).strip())
    except Exception as e:
        print(f"Error extracting selectable text from PDF: {str(e)}")

    # 2) Visual analysis per page (limited)
    try:
        pdf_doc = pdfium.PdfDocument(file_path)
        page_count = len(pdf_doc)
        pages_to_process = min(page_count, max_pages_for_vision)

        for idx in range(pages_to_process):
            try:
                page = pdf_doc[idx]
                # Render page to bitmap and convert to PIL image
                bitmap = page.render(scale=render_scale)
                pil_image = bitmap.to_pil()

                # Save to a temp file for reuse with existing analyzer
                with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
                    temp_path = tmp.name
                    pil_image.save(temp_path, format="PNG")

                try:
                    analysis = analyze_image_with_vision_model(temp_path)
                    text_sections.append(f"Visual analysis for page {idx+1} (AI vision):\n{analysis}".strip())
                finally:
                    try:
                        os.remove(temp_path)
                    except Exception:
                        pass
            except Exception as page_err:
                print(f"Error analyzing page {idx+1} with vision model: {str(page_err)}")
    except Exception as e:
        print(f"Error during visual PDF processing: {str(e)}")

    combined = "\n\n---\n\n".join([s for s in text_sections if s])
    return combined.strip()

# Function to extract text from images using PyTesseract OCR
def extract_text_from_image(file_path):
    image = Image.open(file_path)
    text = pytesseract.image_to_string(image).strip()
    return text

# Function to extract text from DOCX files
def extract_text_from_docx(file_path):
    try:
        doc = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text.strip()
    except Exception as e:
        print(f"Error extracting text from DOCX: {str(e)}")
        return f"Error extracting text from DOCX: {str(e)}"

# Main function to extract text from a file
def extract_text_from_file(file_path):
    try:
        # Check if file exists
        if not os.path.exists(file_path):
            print(f"Error: File not found at path {file_path}")
            return "Error: File not found"
            
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            try:
                return extract_text_from_pdf(file_path)
            except Exception as e:
                print(f"Error extracting text from PDF: {str(e)}")
                return f"Error extracting text from PDF: {str(e)}"
        elif file_extension in ['.png', '.jpg', '.jpeg']:
            try:
                # Always use the vision model for image processing
                return analyze_image_with_vision_model(file_path)
            except Exception as e:
                print(f"Error extracting text from image: {str(e)}")
                return f"Error extracting text from image: {str(e)}"
        elif file_extension == '.txt':
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    return file.read()
            except Exception as e:
                print(f"Error reading text file: {str(e)}")
                return f"Error reading text file: {str(e)}"
        elif file_extension in ['.docx', '.doc']:
            try:
                return extract_text_from_docx(file_path)
            except Exception as e:
                print(f"Error extracting text from DOCX: {str(e)}")
                return f"Error extracting text from DOCX: {str(e)}"
        else:
            return "Unsupported file format"
    except Exception as e:
        print(f"Unexpected error in extract_text_from_file: {str(e)}")
        return f"Error processing file: {str(e)}"

def detect_if_diagram(image_path):
    """Detect if the image is likely a diagram based on edge detection"""
    image = cv2.imread(image_path)
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    blurred = cv2.GaussianBlur(gray, (5, 5), 0)
    edges = cv2.Canny(blurred, 50, 150)
    
    # Count edges and determine if it's likely a diagram
    edge_count = np.sum(edges > 0)
    total_pixels = edges.shape[0] * edges.shape[1]
    edge_ratio = edge_count / total_pixels
    
    # Diagrams typically have more edges than regular text/photos
    return edge_ratio > 0.05

VISION_MODEL = os.environ.get('VISION_MODEL', 'granite3.2-vision:latest')


def analyze_image_with_vision_model(image_path):
    """Analyze image with vision-capable LLM with preprocessing and retry."""
    try:
        # Preprocess (handles large images, non-RGB, and optional SVG rasterization)
        preprocessed_path = _preprocess_image_for_vision(image_path)

        def _call_vision(encoded_img: str):
            return ollama.chat(
                model=VISION_MODEL,
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert OCR assistant that can extract and analyze text from images. Extract all visible text, explain diagrams, and organize the content in a structured way."
                    },
                    {
                        "role": "user",
                        "content": "Please extract and analyze all text and visual elements from this image. If it contains diagrams, charts, or other visual elements, please describe them in detail as well.",
                        "images": [encoded_img]
                    }
                ]
            )

        # Read and encode
        with open(preprocessed_path, "rb") as f:
            image_bytes = f.read()
        encoded_image = base64.b64encode(image_bytes).decode("utf-8")

        print(f"Using {VISION_MODEL} model for image analysis...")

        try:
            response = _call_vision(encoded_image)
            print("Successfully processed image with vision model")
            return response['message']['content']
        except Exception as first_err:
            # Retry once with stronger downscale/re-encode if model complains about embeddings/format
            msg = str(first_err)
            print(f"Error with vision model: {msg}. Retrying with stronger preprocessing...")
            try:
                retry_path = _preprocess_image_for_vision(preprocessed_path, force_small=True)
                with open(retry_path, "rb") as f2:
                    retry_bytes = f2.read()
                encoded_retry = base64.b64encode(retry_bytes).decode("utf-8")
                response = _call_vision(encoded_retry)
                print("Successfully processed image with vision model on retry")
                return response['message']['content']
            except Exception as retry_err:
                print(f"Retry failed: {str(retry_err)}. Falling back to OCR...")
                return extract_text_from_image(image_path)
            finally:
                try:
                    if retry_path != image_path and os.path.exists(retry_path) and retry_path != preprocessed_path:
                        os.remove(retry_path)
                except Exception:
                    pass
        finally:
            try:
                if preprocessed_path != image_path and os.path.exists(preprocessed_path):
                    os.remove(preprocessed_path)
            except Exception:
                pass
    except Exception as e:
        print(f"Error analyzing image: {str(e)}")
        return f"Error analyzing image: {str(e)}"


def _preprocess_image_for_vision(image_path: str, force_small: bool = False) -> str:
    """Prepare image for vision model: handle SVG, ensure RGB, downscale, encode PNG.

    Returns path to a temporary PNG file.
    """
    try:
        suffix = os.path.splitext(image_path)[1].lower()
        temp_path = None

        # Rasterize SVG if cairosvg is available
        if suffix == '.svg':
            try:
                import cairosvg  # type: ignore
                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                    temp_path = tmp.name
                cairosvg.svg2png(url=image_path, write_to=temp_path)
                image_path = temp_path
            except Exception:
                # Fall back: skip SVG rasterization and let OCR handle if possible
                pass

        # Load via PIL
        with Image.open(image_path) as img:
            # Convert to RGB to avoid palette/alpha issues
            if img.mode not in ('RGB', 'L'):
                img = img.convert('RGB')
            # Downscale large images to reduce payload/embedding issues
            max_side = 1600 if not force_small else 1024
            w, h = img.size
            scale = min(1.0, float(max_side) / float(max(w, h) or 1))
            if scale < 1.0:
                new_size = (max(1, int(w * scale)), max(1, int(h * scale)))
                img = img.resize(new_size, Image.LANCZOS)

            # Save to temp PNG with reasonable compression
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as out:
                out_path = out.name
            img.save(out_path, format='PNG', optimize=True)
            return out_path
    except Exception:
        # If preprocessing fails, return original path
        return image_path


def extract_text_with_metadata(file_path: str, mode: str = 'auto'):
    """Extract text and return metadata about the method and processing.

    mode: 'auto' | 'vision_only' | 'ocr_only'
    - vision_only: force granite vision for images; PDFs still combine text + limited vision
    - ocr_only: avoid vision model; use Tesseract for images; PDFs use pdfplumber only
    """
    meta = {
        'success': False,
        'mode': mode,
        'method': None,
        'warnings': [],
        'fileExtension': None,
        'pagesProcessed': 0,
        'visionPages': 0,
    }

    try:
        if not os.path.exists(file_path):
            meta['warnings'].append('File not found')
            return {'text': '', 'meta': meta}

        file_extension = os.path.splitext(file_path)[1].lower()
        meta['fileExtension'] = file_extension

        # PDFs
        if file_extension == '.pdf':
            text_sections = []
            if mode == 'vision_only':
                # Only run vision on rendered pages, skip pdfplumber entirely
                try:
                    pdf_doc = pdfium.PdfDocument(file_path)
                    meta['pagesProcessed'] = len(pdf_doc)
                    pages_to_process = min(len(pdf_doc), int(os.environ.get('VISION_PDF_PAGES', '5')))
                    meta['visionPages'] = pages_to_process
                    for idx in range(pages_to_process):
                        try:
                            page = pdf_doc[idx]
                            bitmap = page.render(scale=2.0)
                            pil_image = bitmap.to_pil()
                            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                                temp_path = tmp.name
                                pil_image.save(temp_path, format='PNG')
                            try:
                                analysis = analyze_image_with_vision_model(temp_path)
                                text_sections.append(f"Visual analysis for page {idx+1} (AI vision):\n{analysis}".strip())
                                meta['method'] = 'vision'
                            finally:
                                try:
                                    os.remove(temp_path)
                                except Exception:
                                    pass
                        except Exception as page_err:
                            meta['warnings'].append(f'Vision page {idx+1} failed: {str(page_err)}')
                except Exception as e:
                    meta['warnings'].append(f'PDF vision processing failed: {str(e)}')
            else:
                # Selectable text via pdfplumber
                try:
                    with pdfplumber.open(file_path) as pdf:
                        meta['pagesProcessed'] = len(pdf.pages)
                        extracted = []
                        for page in pdf.pages:
                            page_text = page.extract_text() or ''
                            if page_text.strip():
                                extracted.append(page_text)
                        if extracted:
                            text_sections.append("Selectable text from PDF:\n" + "\n\n".join(extracted).strip())
                            if meta['method'] is None:
                                meta['method'] = 'pdfplumber'
                except Exception as e:
                    meta['warnings'].append(f'pdfplumber failed: {str(e)}')

                # Vision augmentation for first few pages when not ocr_only
                if mode != 'ocr_only':
                    try:
                        pdf_doc = pdfium.PdfDocument(file_path)
                        pages_to_process = min(len(pdf_doc), int(os.environ.get('VISION_PDF_PAGES', '5')))
                        meta['visionPages'] = pages_to_process
                        for idx in range(pages_to_process):
                            try:
                                page = pdf_doc[idx]
                                bitmap = page.render(scale=2.0)
                                pil_image = bitmap.to_pil()
                                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                                    temp_path = tmp.name
                                    pil_image.save(temp_path, format='PNG')
                                try:
                                    analysis = analyze_image_with_vision_model(temp_path)
                                    text_sections.append(f"Visual analysis for page {idx+1} (AI vision):\n{analysis}".strip())
                                    meta['method'] = (meta['method'] + '+vision') if meta['method'] else 'vision'
                                finally:
                                    try:
                                        os.remove(temp_path)
                                    except Exception:
                                        pass
                            except Exception as page_err:
                                meta['warnings'].append(f'Vision page {idx+1} failed: {str(page_err)}')
                    except Exception as e:
                        meta['warnings'].append(f'PDF vision augmentation failed: {str(e)}')

            text = ("\n\n---\n\n").join([s for s in text_sections if s]).strip()
            meta['success'] = bool(text)
            return {'text': text, 'meta': meta}

        # Images
        if file_extension in ['.png', '.jpg', '.jpeg']:
            if mode == 'ocr_only':
                try:
                    text = extract_text_from_image(file_path)
                    meta['method'] = 'tesseract'
                    meta['success'] = bool(text.strip())
                    return {'text': text, 'meta': meta}
                except Exception as e:
                    meta['warnings'].append(f'OCR failed: {str(e)}')
                    return {'text': '', 'meta': meta}
            else:
                # vision_only or auto -> try vision first, fallback to OCR
                try:
                    text = analyze_image_with_vision_model(file_path)
                    meta['method'] = 'vision'
                    meta['success'] = bool(text.strip())
                    return {'text': text, 'meta': meta}
                except Exception as e:
                    meta['warnings'].append(f'Vision failed: {str(e)}')
                    try:
                        text = extract_text_from_image(file_path)
                        meta['method'] = 'tesseract'
                        meta['success'] = bool(text.strip())
                        return {'text': text, 'meta': meta}
                    except Exception as e2:
                        meta['warnings'].append(f'OCR fallback failed: {str(e2)}')
                        return {'text': '', 'meta': meta}

        # Text
        if file_extension == '.txt':
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                meta['method'] = 'text'
                meta['success'] = True
                return {'text': text, 'meta': meta}
            except Exception as e:
                meta['warnings'].append(f'TXT read failed: {str(e)}')
                return {'text': '', 'meta': meta}

        # Word
        if file_extension in ['.docx', '.doc']:
            text = extract_text_from_docx(file_path)
            meta['method'] = 'docx'
            meta['success'] = bool(text.strip())
            return {'text': text, 'meta': meta}

        meta['warnings'].append('Unsupported file format')
        return {'text': '', 'meta': meta}
    except Exception as e:
        meta['warnings'].append(f'Unexpected error: {str(e)}')
        return {'text': '', 'meta': meta}

def process_ocr_question(content, question, lang: str = 'en'):
    """Process a question based on OCR content"""
    language_hint = f"Please answer in {lang}." if lang and lang.lower() != 'en' else ""
    prompt = f"""Based on the following extracted content:
    
    {content}
    
    Answer this question: {question}
    
    Give a clear, educational answer using only information from the content.
    If the information is not available in the content, state that clearly.
    {language_hint}
    """
    
    try:
        response = ollama.chat(model="mistral", messages=[{"role": "user", "content": prompt}])
        return response['message']['content']
    except Exception as e:
        return f"Error processing question: {str(e)}" 